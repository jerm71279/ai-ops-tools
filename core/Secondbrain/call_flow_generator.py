#!/usr/bin/env python3
"""
Call Flow Generator - OberaConnect
Creates and manages standardized customer call flow configurations

Based on 20+ call flow documents found in SharePoint
Standardizes phone system setup for OberaConnect customers
"""

from typing import Dict, Any, List, Optional
from pathlib import Path
from datetime import datetime
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class CallFlowGenerator:
    """Generate standardized call flow documents for OberaConnect customers"""

    TEMPLATE_FIELDS = {
        "business_info": [
            "business_name",
            "address",
            "contact_name",
            "contact_number",
            "business_hours",
            "main_phone_number"
        ],
        "current_providers": [
            "internet_provider",
            "it_vendor",
            "phone_vendor"
        ],
        "phone_setup": [
            "creating_or_porting",  # "porting" or "creating"
            "numbers_to_port",  # List of phone numbers if porting
            "new_numbers_needed",  # Count and area code if creating
            "requested_port_date"
        ],
        "configuration": [
            "extension_list",  # List of extensions with names/emails
            "blf_keys",  # Button configuration
            "portal_access",  # Who needs ConnectWare access
            "mobile_app_users",  # ConnectMobile users
            "texting_type",  # internal/external
            "texting_numbers",  # Which numbers need texting
        ],
        "call_routing": [
            "business_hours_routing",
            "auto_attendant_script",
            "after_hours_routing",
            "after_hours_script",
            "after_hours_vm_message",
            "voicemail_type",  # standard or VM-to-email
            "holiday_schedule"
        ],
        "equipment": [
            "phone_count",
            "cordless_count",
            "conference_phone_count",
            "wall_mounted_count",
            "headsets",
            "fax_setup"  # EFax or Analog
        ],
        "additional": [
            "phone_training_needed",
            "integrations",  # CRM or other software
            "paging_setup",
            "alarm_lines",
            "elevator_lines",
            "call_recording",
            "reporting_needs"
        ]
    }

    def __init__(self):
        self.output_dir = Path("call_flows_generated")
        self.output_dir.mkdir(exist_ok=True)
        print(f"✓ Call Flow Generator initialized")
        print(f"  Output directory: {self.output_dir}")

    def create_call_flow(self, customer_data: Dict[str, Any]) -> Path:
        """
        Create a new call flow document from customer data

        Args:
            customer_data: Dictionary with call flow information

        Returns:
            Path to generated DOCX file
        """

        business_name = customer_data.get("business_name", "Unknown Customer")

        # Create document
        doc = Document()

        # Title
        title = doc.add_heading(f'Call Flow Configuration', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_heading(business_name, level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Date
        date_para = doc.add_paragraph(f'Created: {datetime.now().strftime("%Y-%m-%d")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spacing

        # Business Information Section
        self._add_section(doc, "Business Information", [
            ("Business Name", customer_data.get("business_name", "")),
            ("Address", customer_data.get("address", "")),
            ("Contact Name", customer_data.get("contact_name", "")),
            ("Contact Number", customer_data.get("contact_number", "")),
            ("Business Hours", customer_data.get("business_hours", "")),
            ("Main Phone Number", customer_data.get("main_phone_number", ""))
        ])

        # Current Providers Section
        self._add_section(doc, "Current Service Providers", [
            ("Internet Service Provider", customer_data.get("internet_provider", "")),
            ("Current IT Vendor", customer_data.get("it_vendor", "")),
            ("Current Phone Vendor", customer_data.get("phone_vendor", ""))
        ])

        # Phone Number Setup
        creating_or_porting = customer_data.get("creating_or_porting", "")
        phone_fields = [
            ("Creating or Porting", creating_or_porting)
        ]

        if creating_or_porting.lower() == "porting":
            numbers = customer_data.get("numbers_to_port", [])
            if isinstance(numbers, list):
                numbers_str = ", ".join(numbers)
            else:
                numbers_str = str(numbers)
            phone_fields.append(("Numbers to Port", numbers_str))
            phone_fields.append(("⚠️ IMPORTANT", "Notify customer when they can cancel services. Early cancellation = lost numbers!"))
        else:
            phone_fields.append(("New Numbers Needed", customer_data.get("new_numbers_needed", "")))

        phone_fields.append(("Requested Port Date", customer_data.get("requested_port_date", "")))

        self._add_section(doc, "Phone Number Setup", phone_fields)

        # Configuration
        self._add_section(doc, "System Configuration", [
            ("Extension/Configuration List", customer_data.get("extension_list", "")),
            ("BLF Keys (External)", customer_data.get("blf_keys", "Internal already programmed")),
            ("ConnectWare Portal Access", customer_data.get("portal_access", "")),
            ("ConnectMobile App Users", customer_data.get("mobile_app_users", "")),
            ("Texting Type", customer_data.get("texting_type", "")),
            ("Texting-Enabled Numbers", customer_data.get("texting_numbers", "")),
        ])

        # Call Routing
        self._add_section(doc, "Call Routing", [
            ("Business Hours Routing", customer_data.get("business_hours_routing", "")),
            ("Auto Attendant Script", customer_data.get("auto_attendant_script", "")),
            ("After Hours Routing", customer_data.get("after_hours_routing", "")),
            ("After Hours AA Script", customer_data.get("after_hours_script", "")),
            ("After Hours VM Message", customer_data.get("after_hours_vm_message", "")),
            ("Voicemail Type", customer_data.get("voicemail_type", "VM-to-email")),
            ("Holiday Schedule", customer_data.get("holiday_schedule", "Email scripts to support@oberaconnect.com"))
        ])

        # Equipment
        self._add_section(doc, "Equipment Requirements", [
            ("Phone Count", customer_data.get("phone_count", "")),
            ("Cordless Phone Count", customer_data.get("cordless_count", "")),
            ("Conference Phone Count", customer_data.get("conference_phone_count", "")),
            ("Wall Mounted Phones", customer_data.get("wall_mounted_count", "")),
            ("Headsets", customer_data.get("headsets", "")),
            ("Fax Setup", customer_data.get("fax_setup", ""))
        ])

        # Additional Services
        self._add_section(doc, "Additional Services", [
            ("Phone Training", customer_data.get("phone_training_needed", "")),
            ("CRM/Software Integrations", customer_data.get("integrations", "")),
            ("Paging Setup", customer_data.get("paging_setup", "")),
            ("Alarm Lines", customer_data.get("alarm_lines", "")),
            ("Elevator Lines", customer_data.get("elevator_lines", "")),
            ("Call Recording", customer_data.get("call_recording", "")),
            ("Reporting Needs", customer_data.get("reporting_needs", ""))
        ])

        # Installation Notes
        doc.add_heading("Installation Notes", level=2)
        notes = doc.add_paragraph()
        notes.add_run("Installation Process:\n").bold = True
        notes.add_run("• Confirm fiber availability and installation timeline\n")
        notes.add_run("• Notify customer of installation date\n")
        notes.add_run("• Verify equipment delivery before install day\n")
        notes.add_run("• Notify customer when services can be cancelled\n")

        # Save document
        filename = f"{business_name.replace(' ', '_')}_CallFlow_{datetime.now().strftime('%Y%m%d')}.docx"
        output_path = self.output_dir / filename
        doc.save(output_path)

        print(f"✓ Call flow generated: {output_path}")
        return output_path

    def _add_section(self, doc: Document, heading: str, fields: List[tuple]):
        """Add a formatted section to the document"""
        doc.add_heading(heading, level=2)

        for label, value in fields:
            para = doc.add_paragraph()
            para.add_run(f"{label}: ").bold = True
            para.add_run(str(value))

        doc.add_paragraph()  # Spacing

    def create_template(self) -> Dict[str, Any]:
        """Return empty template structure for creating new call flows"""
        template = {}
        for category, fields in self.TEMPLATE_FIELDS.items():
            for field in fields:
                template[field] = ""
        return template

    def export_to_json(self, customer_data: Dict[str, Any], filename: str = None):
        """Export call flow data to JSON for storage"""
        if not filename:
            business_name = customer_data.get("business_name", "unknown")
            filename = f"{business_name.replace(' ', '_')}_callflow.json"

        output_path = self.output_dir / filename
        with open(output_path, 'w') as f:
            json.dump(customer_data, f, indent=2)

        print(f"✓ Call flow data exported: {output_path}")
        return output_path

    def import_from_json(self, json_path: Path) -> Dict[str, Any]:
        """Import call flow data from JSON"""
        with open(json_path, 'r') as f:
            return json.load(f)


def main():
    """Example usage"""
    generator = CallFlowGenerator()

    # Example: Create a call flow from template
    print("\n📋 Creating example call flow...")

    example_data = {
        "business_name": "Example Business Inc",
        "address": "123 Main St, Anytown, FL 32561",
        "contact_name": "John Smith",
        "contact_number": "850-555-1234",
        "business_hours": "Mon-Fri 8:00 AM - 5:00 PM",
        "main_phone_number": "850-555-1000",
        "internet_provider": "Uniti",
        "it_vendor": "OberaConnect",
        "phone_vendor": "Nextiva",
        "creating_or_porting": "porting",
        "numbers_to_port": ["8505551000", "8505551001"],
        "requested_port_date": "2025-12-01",
        "extension_list": "Reception: ext 100, Sales: ext 101",
        "portal_access": "John Smith, Jane Doe",
        "mobile_app_users": "John Smith",
        "texting_type": "external",
        "texting_numbers": "8505551000",
        "business_hours_routing": "Ring all phones",
        "after_hours_routing": "Voicemail",
        "voicemail_type": "VM-to-email",
        "phone_count": "5",
        "cordless_count": "2",
        "phone_training_needed": "Portal training"
    }

    # Generate document
    output_file = generator.create_call_flow(example_data)

    # Also save as JSON
    json_file = generator.export_to_json(example_data)

    print(f"\n✅ Generated:")
    print(f"   DOCX: {output_file}")
    print(f"   JSON: {json_file}")

    # Show template structure
    print("\n📄 Template structure available:")
    template = generator.create_template()
    for field in sorted(template.keys()):
        print(f"   - {field}")


if __name__ == "__main__":
    main()
