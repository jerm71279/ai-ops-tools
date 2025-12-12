#!/usr/bin/env python3
"""
SOP Markdown to Scribe-Style DOCX Converter

Converts SOP markdown files to professionally formatted DOCX files
matching the Scribe-Style SOP Template format.
"""

import os
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# Color definitions for status badges
STATUS_COLORS = {
    'DRAFT': {'bg': 'FFC7CE', 'text': '9C0006'},
    'IN REVIEW': {'bg': 'FFEB9C', 'text': '9C5700'},
    'APPROVED': {'bg': 'C6EFCE', 'text': '006100'},
    'PUBLISHED': {'bg': 'BDD7EE', 'text': '1F4E79'},
    'ARCHIVED': {'bg': 'E2EFDA', 'text': '375623'},
    'DEPRECATED': {'bg': 'D9D9D9', 'text': '595959'},
}

HEADER_BG = 'D9E2F3'  # Light blue for table headers
LABEL_BG = 'F2F2F2'   # Light gray for label cells


def set_cell_shading(cell, fill_color):
    """Set background color for a table cell."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{fill_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_border(cell, border_color='CCCCCC', border_size='4'):
    """Set borders for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'''<w:tcBorders {nsdecls("w")}>
            <w:top w:val="single" w:color="{border_color}" w:sz="{border_size}"/>
            <w:left w:val="single" w:color="{border_color}" w:sz="{border_size}"/>
            <w:bottom w:val="single" w:color="{border_color}" w:sz="{border_size}"/>
            <w:right w:val="single" w:color="{border_color}" w:sz="{border_size}"/>
        </w:tcBorders>'''
    )
    tcPr.append(tcBorders)


def add_formatted_table(doc, headers, rows, col_widths=None, header_bg=HEADER_BG):
    """Create a formatted table with headers and rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.bold = True
        run.font.size = Pt(10)
        set_cell_shading(cell, header_bg)
        set_cell_border(cell)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            cell = row_cells[col_idx]
            cell.text = str(cell_data) if cell_data else ''
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(10)
            set_cell_border(cell)

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    return table


def add_two_column_table(doc, data, label_width=2.0, value_width=5.0):
    """Create a two-column label/value table."""
    table = doc.add_table(rows=len(data), cols=2)

    for row_idx, (label, value) in enumerate(data):
        label_cell = table.rows[row_idx].cells[0]
        value_cell = table.rows[row_idx].cells[1]

        label_cell.text = label
        value_cell.text = value

        # Style label cell
        para = label_cell.paragraphs[0]
        run = para.runs[0]
        run.bold = True
        run.font.size = Pt(10)
        set_cell_shading(label_cell, LABEL_BG)
        set_cell_border(label_cell)

        # Style value cell
        para = value_cell.paragraphs[0]
        for run in para.runs:
            run.font.size = Pt(10)
        set_cell_border(value_cell)

        label_cell.width = Inches(label_width)
        value_cell.width = Inches(value_width)

    return table


def add_status_badge_row(doc):
    """Add the status badge row with colored cells."""
    statuses = ['DRAFT', 'IN REVIEW', 'APPROVED', 'PUBLISHED', 'ARCHIVED', 'DEPRECATED']
    table = doc.add_table(rows=1, cols=len(statuses))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, status in enumerate(statuses):
        cell = table.rows[0].cells[i]
        cell.text = status
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.bold = True
        run.font.size = Pt(8)
        colors = STATUS_COLORS[status]
        run.font.color.rgb = RGBColor.from_string(colors['text'])
        set_cell_shading(cell, colors['bg'])
        set_cell_border(cell)

    return table


def add_checkbox_row(doc, status='APPROVED', classification='Internal', review_cycle='Semi-Annual'):
    """Add status/classification/review checkbox row."""
    status_opts = ['Draft', 'In Review', 'Approved', 'Retired']
    class_opts = ['Public', 'Internal', 'Confidential']
    review_opts = ['Quarterly', 'Semi-Annual', 'Annual']

    def format_checkboxes(options, selected):
        parts = []
        for opt in options:
            check = '☑' if opt.lower() == selected.lower() else '☐'
            parts.append(f'{check} {opt}')
        return '  '.join(parts)

    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Headers
    headers = ['Status', 'Classification', 'Review Cycle']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.bold = True
        run.font.size = Pt(10)
        set_cell_shading(cell, HEADER_BG)
        set_cell_border(cell)

    # Checkbox values
    values = [
        format_checkboxes(status_opts, status.replace('IN ', 'In ')),
        format_checkboxes(class_opts, classification),
        format_checkboxes(review_opts, review_cycle)
    ]
    for i, value in enumerate(values):
        cell = table.rows[1].cells[i]
        cell.text = value
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.size = Pt(9)
        set_cell_border(cell)

    return table


def add_verification_checklist(doc, items):
    """Add verification checklist with checkboxes."""
    headers = ['✓', 'Verification Item', 'Verified By / Date']
    rows = [['☐', item, ''] for item in items]
    return add_formatted_table(doc, headers, rows, col_widths=[0.3, 5.0, 1.7])


def add_troubleshooting_table(doc, issues):
    """Add troubleshooting table."""
    headers = ['Issue / Symptom', 'Possible Cause', 'Resolution']
    return add_formatted_table(doc, headers, issues, col_widths=[2.0, 2.0, 3.0])


def add_revision_history_table(doc, revisions):
    """Add revision history table."""
    headers = ['Version', 'Date', 'Status', 'Author', 'Change Description']
    return add_formatted_table(doc, headers, revisions, col_widths=[0.8, 1.0, 1.0, 1.0, 3.2])


def add_related_documents_table(doc, documents):
    """Add related documents table."""
    headers = ['Doc ID', 'Document Title', 'Relationship']
    return add_formatted_table(doc, headers, documents, col_widths=[1.5, 3.5, 2.0])


def add_roles_table(doc, roles):
    """Add roles and responsibilities table."""
    headers = ['Role', 'Name/Position', 'Responsibilities']
    return add_formatted_table(doc, headers, roles, col_widths=[1.5, 1.5, 4.0])


def add_definitions_table(doc, definitions):
    """Add definitions/acronyms table."""
    headers = ['Term / Acronym', 'Definition']
    return add_formatted_table(doc, headers, definitions, col_widths=[2.0, 5.0])


def add_approval_table(doc):
    """Add approval signature table."""
    table = doc.add_table(rows=2, cols=2)

    data = [
        ('Process Owner', 'Name:\nSignature:\nDate:'),
        ('Department Head', 'Name:\nSignature:\nDate:')
    ]

    for row_idx, (label, value) in enumerate(data):
        label_cell = table.rows[row_idx].cells[0]
        value_cell = table.rows[row_idx].cells[1]

        label_cell.text = label
        value_cell.text = value

        para = label_cell.paragraphs[0]
        run = para.runs[0]
        run.bold = True
        run.font.size = Pt(10)
        set_cell_shading(label_cell, LABEL_BG)
        set_cell_border(label_cell)
        set_cell_border(value_cell)

    return table


def parse_markdown_sop(md_content):
    """Parse markdown SOP content into structured data."""
    # Remove leading "Loaded cached credentials" line if present
    md_content = re.sub(r'^Loaded cached credentials\.\n', '', md_content)
    # Remove markdown code block wrapper if present
    md_content = re.sub(r'^```markdown\n', '', md_content)
    md_content = re.sub(r'\n```$', '', md_content)

    data = {
        'title': '',
        'doc_id': '',
        'version': '1.0.0',
        'status': 'APPROVED',
        'author': '',
        'category': '',
        'creation_date': '',
        'approval_date': '',
        'purpose': '',
        'scope': '',
        'definitions': [],
        'roles': [],
        'prerequisites': [],
        'procedure_sections': [],
        'verification': [],
        'troubleshooting': [],
        'related_docs': [],
        'revision_history': [],
    }

    # Extract title
    title_match = re.search(r'^#\s+Standard Operating Procedure:\s*(.+)$', md_content, re.MULTILINE)
    if title_match:
        data['title'] = title_match.group(1).strip()

    # Extract metadata from table
    doc_id_match = re.search(r'\*\*Document ID:\*\*\s*\|\s*(\S+)', md_content)
    if doc_id_match:
        data['doc_id'] = doc_id_match.group(1)

    version_match = re.search(r'\*\*Version:\*\*\s*\|\s*(\S+)', md_content)
    if version_match:
        v = version_match.group(1)
        data['version'] = v if '.' in v and v.count('.') == 2 else f"{v}.0"

    status_match = re.search(r'\*\*Status:\*\*\s*\|\s*(\S+)', md_content)
    if status_match:
        data['status'] = status_match.group(1).upper()

    author_match = re.search(r'\*\*Author:\*\*\s*\|\s*(.+)', md_content)
    if author_match:
        data['author'] = author_match.group(1).strip()

    category_match = re.search(r'\*\*Category:\*\*\s*\|\s*(.+)', md_content)
    if category_match:
        data['category'] = category_match.group(1).strip()

    creation_match = re.search(r'\*\*Creation Date:\*\*\s*\|\s*(\S+)', md_content)
    if creation_match:
        data['creation_date'] = creation_match.group(1)

    approval_match = re.search(r'\*\*Approval Date:\*\*\s*\|\s*(\S+)', md_content)
    if approval_match:
        data['approval_date'] = approval_match.group(1)

    # Extract sections
    def extract_section(pattern, content):
        match = re.search(pattern, content, re.DOTALL)
        return match.group(1).strip() if match else ''

    data['purpose'] = extract_section(r'###\s*1\.0\s+Purpose\s*\n(.+?)(?=###|\Z)', md_content)
    data['scope'] = extract_section(r'###\s*2\.0\s+Scope\s*\n(.+?)(?=###|\Z)', md_content)

    # Extract definitions
    definitions_section = extract_section(r'###\s*3\.0\s+Definitions\s*\n(.+?)(?=###|\Z)', md_content)
    if definitions_section:
        for line in definitions_section.split('\n'):
            match = re.match(r'\*\s+\*\*(.+?):\*\*\s*(.+)', line)
            if match:
                data['definitions'].append([match.group(1), match.group(2)])

    # Extract roles
    roles_section = extract_section(r'###\s*4\.0\s+Roles.*?\n(.+?)(?=###|\Z)', md_content)
    if roles_section:
        for line in roles_section.split('\n'):
            match = re.match(r'\*\s+\*\*(.+?):\*\*\s*(.+)', line)
            if match:
                data['roles'].append([match.group(1), '', match.group(2)])

    # Extract prerequisites
    prereq_section = extract_section(r'###\s*5\.0\s+Prerequisites\s*\n(.+?)(?=---|\Z)', md_content)
    if prereq_section:
        for line in prereq_section.split('\n'):
            if line.strip().startswith('*'):
                item = re.sub(r'^\*\s+', '', line.strip())
                data['prerequisites'].append(item)

    # Extract procedure sections
    procedure_section = extract_section(r'###\s*6\.0\s+Procedure\s*\n(.+?)(?=###\s*7\.0|\Z)', md_content)
    if procedure_section:
        # Split into subsections
        subsections = re.split(r'####\s*(\d+\.\d+)\s+(.+?)\n', procedure_section)
        for i in range(1, len(subsections), 3):
            if i + 2 < len(subsections):
                section_num = subsections[i]
                section_title = subsections[i + 1]
                section_content = subsections[i + 2]

                steps = []
                # Parse numbered steps
                step_matches = re.findall(r'(\d+)\.\s+\*\*(.+?):\*\*\s*\n?(.*?)(?=\d+\.\s+\*\*|\Z)',
                                         section_content, re.DOTALL)
                for step_num, step_title, step_content in step_matches:
                    step_content = step_content.strip()
                    # Extract expected result if present
                    expected = ''
                    exp_match = re.search(r'\*\*Expected Result:\*\*\s*(.+?)(?=\n|$)', step_content)
                    if exp_match:
                        expected = exp_match.group(1)

                    # Extract warning if present
                    warning = ''
                    warn_match = re.search(r'>\s*\*\*WARNING:\*\*\s*(.+?)(?=\n|$)', step_content, re.IGNORECASE)
                    if not warn_match:
                        warn_match = re.search(r'>\s*WARNING:\s*(.+?)(?=\n|$)', step_content, re.IGNORECASE)
                    if warn_match:
                        warning = warn_match.group(1)

                    steps.append({
                        'number': step_num,
                        'title': step_title,
                        'content': step_content,
                        'expected': expected,
                        'warning': warning
                    })

                data['procedure_sections'].append({
                    'number': section_num,
                    'title': section_title,
                    'steps': steps
                })

    # Extract verification
    verif_section = extract_section(r'###\s*7\.0\s+Verification.*?\n(.+?)(?=###|\Z)', md_content)
    if verif_section:
        for line in verif_section.split('\n'):
            if line.strip().startswith('*'):
                item = re.sub(r'^\*\s+', '', line.strip())
                data['verification'].append(item)

    # Extract troubleshooting
    trouble_section = extract_section(r'###\s*8\.0\s+Troubleshooting\s*\n(.+?)(?=###|\Z)', md_content)
    if trouble_section:
        # Parse table rows
        rows = re.findall(r'\|\s*(.+?)\s*\|\s*(.+?)\s*\|', trouble_section)
        for row in rows:
            if row[0] and not row[0].startswith('---') and row[0] != 'Issue':
                parts = [row[0], '', row[1]] if len(row) == 2 else list(row)[:3]
                data['troubleshooting'].append(parts)

    # Extract related documents
    related_section = extract_section(r'###\s*9\.0\s+Related Documents\s*\n(.+?)(?=###|\Z)', md_content)
    if related_section:
        for line in related_section.split('\n'):
            if line.strip().startswith('*'):
                item = re.sub(r'^\*\s+', '', line.strip())
                data['related_docs'].append(['', item, 'Reference'])

    # Extract revision history
    revision_section = extract_section(r'###\s*10\.0\s+Revision History\s*\n(.+?)(?=###|\Z)', md_content)
    if revision_section:
        rows = re.findall(r'\|\s*(.+?)\s*\|\s*(.+?)\s*\|\s*(.+?)\s*\|\s*(.+?)\s*\|', revision_section)
        for row in rows:
            if row[0] and not row[0].startswith('---') and row[0] != 'Version':
                data['revision_history'].append([row[0], row[1], 'APPROVED', row[2], row[3]])

    return data


def create_scribe_style_sop(data, output_path):
    """Create a Scribe-style DOCX SOP from parsed data."""
    doc = Document()

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('STANDARD OPERATING PROCEDURE')
    title_run.bold = True
    title_run.font.size = Pt(24)

    # Subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run(data['title'])
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    doc.add_paragraph()

    # Document metadata table
    effective_date = data['approval_date'] or datetime.now().strftime('%m/%d/%Y')
    # Calculate next review (6 months from effective)
    try:
        eff_dt = datetime.strptime(effective_date, '%Y-%m-%d')
        next_review = eff_dt.replace(year=eff_dt.year + 1).strftime('%m/%d/%Y')
        effective_date = eff_dt.strftime('%m/%d/%Y')
    except:
        next_review = '[MM/DD/YYYY]'

    headers = ['Document ID', 'Version', 'Status', 'Effective Date', 'Next Review']
    status_text = data['status'] if data['status'] in STATUS_COLORS else 'APPROVED'
    rows = [[data['doc_id'], data['version'], status_text, effective_date, next_review]]

    meta_table = add_formatted_table(doc, headers, rows)
    # Color the status cell
    status_cell = meta_table.rows[1].cells[2]
    colors = STATUS_COLORS.get(status_text, STATUS_COLORS['APPROVED'])
    set_cell_shading(status_cell, colors['bg'])
    for para in status_cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(colors['text'])

    doc.add_paragraph()

    # Status badge row
    add_status_badge_row(doc)

    doc.add_paragraph()

    # Status/Classification/Review checkbox row
    add_checkbox_row(doc, status='Approved', classification='Internal', review_cycle='Semi-Annual')

    doc.add_paragraph()

    # Process Owner table
    add_two_column_table(doc, [
        ('Process Owner', '[Name / Role / Department]'),
        ('Author', data['author'] or '[Name]'),
        ('Approved By', '[Name / Title]'),
        ('Department/Team', data['category'] or '[e.g., IT Operations, Service Desk, NOC]')
    ])

    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        'Purpose', 'Scope', 'Definitions & Acronyms', 'Roles & Responsibilities',
        'Prerequisites & Requirements', 'Procedure', 'Verification & Quality Checks',
        'Troubleshooting', 'Related Documents', 'Revision History', 'Approval'
    ]
    for i, item in enumerate(toc_items, 1):
        para = doc.add_paragraph(f'{i}. {item}', style='List Number')

    # 1. Purpose
    doc.add_heading('1. Purpose', level=1)
    doc.add_paragraph(data['purpose'] or 'This Standard Operating Procedure (SOP) describes the process for [brief description].')

    # 2. Scope
    doc.add_heading('2. Scope', level=1)
    doc.add_heading('2.1 In Scope', level=2)
    scope_text = data['scope'] or 'This SOP applies to:'
    doc.add_paragraph(scope_text)

    doc.add_heading('2.2 Out of Scope', level=2)
    doc.add_paragraph('This SOP does NOT cover:')
    doc.add_paragraph('[Exclusions]', style='List Bullet')

    # 3. Definitions & Acronyms
    doc.add_heading('3. Definitions & Acronyms', level=1)
    if data['definitions']:
        add_definitions_table(doc, data['definitions'])
    else:
        add_definitions_table(doc, [['[Term]', '[Definition]']])

    # 4. Roles & Responsibilities
    doc.add_heading('4. Roles & Responsibilities', level=1)
    if data['roles']:
        add_roles_table(doc, data['roles'])
    else:
        add_roles_table(doc, [
            ['Process Owner', '[Name/Title]', 'Owns process, approves changes'],
            ['Executor', '[Role]', 'Performs procedure steps'],
            ['Reviewer', '[Name/Title]', 'Reviews work, provides sign-off']
        ])

    # 5. Prerequisites & Requirements
    doc.add_heading('5. Prerequisites & Requirements', level=1)

    doc.add_heading('5.1 Required Access & Permissions', level=2)
    if data['prerequisites']:
        for prereq in data['prerequisites']:
            doc.add_paragraph(prereq, style='List Bullet')
    else:
        doc.add_paragraph('[System access requirements]', style='List Bullet')

    doc.add_heading('5.2 Required Tools & Resources', level=2)
    doc.add_paragraph('[Software/hardware requirements]', style='List Bullet')

    doc.add_heading('5.3 Pre-Conditions', level=2)
    doc.add_paragraph('[State requirements]', style='List Bullet')

    # 6. Procedure
    doc.add_heading('6. Procedure', level=1)
    time_para = doc.add_paragraph()
    time_run = time_para.add_run('⏱ Estimated Time: ')
    time_run.bold = True
    time_para.add_run('[XX minutes/hours]')

    if data['procedure_sections']:
        for section in data['procedure_sections']:
            doc.add_heading(f'{section["number"]} {section["title"]}', level=2)

            for step in section['steps']:
                # Step title
                step_para = doc.add_paragraph()
                step_run = step_para.add_run(f'Step {step["number"]}: {step["title"]}')
                step_run.bold = True

                # Step content
                content = step['content']
                # Clean up content - remove expected result and warning lines
                content = re.sub(r'\*\*Expected Result:\*\*.*', '', content)
                content = re.sub(r'>.*?WARNING.*?\n?', '', content, flags=re.IGNORECASE)
                content = content.strip()
                if content:
                    # Parse bullet points
                    lines = content.split('\n')
                    for line in lines:
                        line = line.strip()
                        if line.startswith('*'):
                            doc.add_paragraph(line[1:].strip(), style='List Bullet')
                        elif line:
                            doc.add_paragraph(line)

                # Expected result
                if step['expected']:
                    exp_para = doc.add_paragraph()
                    exp_run = exp_para.add_run('Expected Result: ')
                    exp_run.bold = True
                    exp_para.add_run(step['expected'])

                # Warning
                if step['warning']:
                    warn_para = doc.add_paragraph()
                    warn_run = warn_para.add_run('⚠️ Warning: ')
                    warn_run.bold = True
                    warn_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                    warn_para.add_run(step['warning'])
    else:
        doc.add_heading('6.1 [Phase/Section Name]', level=2)
        doc.add_paragraph('Step 1: [Action Title]', style='List Number')

    # 7. Verification & Quality Checks
    doc.add_heading('7. Verification & Quality Checks', level=1)
    doc.add_paragraph('Complete the following checklist to verify successful procedure completion:')

    verif_items = data['verification'] if data['verification'] else [
        '[Verification check 1]',
        '[Verification check 2]',
        '[Verification check 3]'
    ]
    add_verification_checklist(doc, verif_items)

    # 8. Troubleshooting
    doc.add_heading('8. Troubleshooting', level=1)
    doc.add_paragraph('Common issues and their resolutions:')

    trouble_items = data['troubleshooting'] if data['troubleshooting'] else [
        ['[Issue description]', '[Root cause]', '[Steps to resolve]']
    ]
    add_troubleshooting_table(doc, trouble_items)

    doc.add_heading('8.1 Escalation Path', level=2)
    doc.add_paragraph('If issues cannot be resolved using this guide:')
    doc.add_paragraph('Document the issue and steps attempted', style='List Number')
    doc.add_paragraph('Escalate to [Tier 2 / Team Lead / Process Owner]', style='List Number')
    doc.add_paragraph('Contact: [Name] - [Email/Phone] - [Hours]', style='List Number')

    # 9. Related Documents
    doc.add_heading('9. Related Documents', level=1)
    related = data['related_docs'] if data['related_docs'] else [
        ['[SOP-XXX-001]', '[Related document title]', 'Reference']
    ]
    add_related_documents_table(doc, related)

    # 10. Revision History
    doc.add_heading('10. Revision History', level=1)
    revisions = data['revision_history'] if data['revision_history'] else [
        ['1.0.0', data['creation_date'] or datetime.now().strftime('%Y-%m-%d'),
         'APPROVED', data['author'] or '[Name]', 'Initial document creation']
    ]
    add_revision_history_table(doc, revisions)

    # 11. Versioning Guidelines
    doc.add_heading('11. Versioning Guidelines', level=1)
    doc.add_paragraph('This document follows semantic versioning (Major.Minor.Patch) to track changes:')

    version_headers = ['Version Type', 'When to Use', 'Example Changes']
    version_rows = [
        ['Major (X.0.0)', 'Complete rewrite, fundamental process change', 'New platform migration, workflow redesign'],
        ['Minor (x.X.0)', 'New steps added, sections reorganized', 'Added verification steps, new section'],
        ['Patch (x.x.X)', 'Typo fixes, clarifications, minor corrections', 'Fixed screenshot, updated contact info']
    ]
    add_formatted_table(doc, version_headers, version_rows)

    doc.add_heading('11.1 Document Status Definitions', level=2)
    status_headers = ['Status', 'Description']
    status_rows = [
        ['Draft', 'Work in progress. Not approved for production use.'],
        ['In Review', 'Submitted for stakeholder review. Pending approval.'],
        ['Approved', 'Officially approved for production use.'],
        ['Retired', 'No longer active. Archived for reference only.']
    ]
    add_formatted_table(doc, status_headers, status_rows)

    # 12. Approval
    doc.add_heading('12. Approval', level=1)
    doc.add_paragraph('By signing below, I confirm this SOP has been reviewed and approved for use.')
    add_approval_table(doc)

    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('This document is the property of OberaConnect. Unauthorized distribution is prohibited.')
    footer_run.italic = True
    footer_run.font.size = Pt(9)

    # Save
    doc.save(output_path)
    return output_path


def convert_sop_file(md_path, output_dir):
    """Convert a single markdown SOP file to DOCX."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    data = parse_markdown_sop(content)

    # Generate output filename
    base_name = Path(md_path).stem
    # Add date to filename
    date_str = datetime.now().strftime('%Y%m%d')
    output_name = f"{base_name}_{date_str}.docx"
    output_path = os.path.join(output_dir, output_name)

    create_scribe_style_sop(data, output_path)
    return output_path


def main():
    """Main entry point."""
    sops_dir = '/home/mavrick/Projects/Secondbrain/SOPs'
    output_dir = '/home/mavrick/Projects/Secondbrain/SOPs/converted_docx'

    # Create output directory
    os.makedirs(output_dir, exist_ok=True)

    # Get all markdown SOP files (exclude README)
    md_files = [f for f in os.listdir(sops_dir)
                if f.endswith('.md') and f.startswith('SOP-')]

    print(f"Found {len(md_files)} SOP files to convert")

    converted = []
    for md_file in sorted(md_files):
        md_path = os.path.join(sops_dir, md_file)
        try:
            output_path = convert_sop_file(md_path, output_dir)
            converted.append(output_path)
            print(f"✓ Converted: {md_file}")
        except Exception as e:
            print(f"✗ Failed: {md_file} - {e}")

    print(f"\nConversion complete: {len(converted)}/{len(md_files)} files")
    print(f"Output directory: {output_dir}")

    return converted


if __name__ == '__main__':
    main()
